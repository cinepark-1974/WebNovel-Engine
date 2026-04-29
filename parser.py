"""
parser.py — 기획서 파서 v3.0 Phase D
=====================================

업로드된 다양한 형식의 기획서 파일에서 텍스트를 추출.
지원 형식: docx, hwp, pdf, txt, md

v3.0 Phase D 보강:
- DOCX: 일반 paragraph + 표 안 paragraph 모두 수집 (이전엔 표 누락)
- DOCX: 헤딩 스타일 보존 (## 마크다운으로 표시)
- 추출 결과의 분량을 호출자가 알 수 있도록 메타정보 함께 반환 옵션
"""

import io


def parse_brief(uploaded_file) -> str:
    """업로드된 파일을 텍스트로 변환.
    
    Args:
        uploaded_file: Streamlit file_uploader에서 받은 UploadedFile 객체
    
    Returns:
        추출된 텍스트.
    """
    if uploaded_file is None:
        return ""
    
    name = uploaded_file.name.lower()
    
    if name.endswith(".docx"):
        return _parse_docx(uploaded_file)
    if name.endswith(".hwp"):
        return _parse_hwp(uploaded_file)
    if name.endswith(".pdf"):
        return _parse_pdf(uploaded_file)
    if name.endswith((".txt", ".md")):
        return _parse_text(uploaded_file)
    
    # 알 수 없는 형식이면 텍스트 시도
    return _parse_text(uploaded_file)


def _parse_docx(uploaded_file) -> str:
    """DOCX 파싱 — 일반 paragraph + 표 안 paragraph + 헤딩 스타일 모두 수집.
    
    [v3.0 Phase D] 이전 버전은 doc.paragraphs만 사용해 표 내부 텍스트가 누락됐다.
    이번 버전은 표 안 paragraph도 같이 수집.
    """
    try:
        from docx import Document
    except ImportError:
        return "[ERROR] python-docx not installed"
    
    try:
        # Streamlit UploadedFile은 BytesIO 호환
        doc = Document(uploaded_file)
    except Exception as e:
        return f"[ERROR] DOCX 읽기 실패: {e}"
    
    parts = []
    
    # body의 element_iter로 paragraph + table을 순서대로 순회
    # python-docx의 doc.element.body 자식들을 순회하면 등장 순서대로 처리 가능
    try:
        body = doc.element.body
        for child in body.iterchildren():
            tag = child.tag.split("}")[-1]  # 네임스페이스 제거
            
            if tag == "p":
                # 일반 paragraph
                from docx.text.paragraph import Paragraph
                para = Paragraph(child, doc)
                text = para.text.strip()
                if not text:
                    continue
                
                # 헤딩 스타일이면 ## 표시 추가 (LLM이 구조 인식하도록)
                style_name = ""
                try:
                    style_name = (para.style.name or "").lower()
                except Exception:
                    pass
                
                if "heading 1" in style_name or "title" in style_name:
                    parts.append(f"# {text}")
                elif "heading 2" in style_name:
                    parts.append(f"## {text}")
                elif "heading 3" in style_name:
                    parts.append(f"### {text}")
                else:
                    parts.append(text)
            
            elif tag == "tbl":
                # 표 — 표 안 paragraph 모두 수집
                from docx.table import Table
                table = Table(child, doc)
                table_lines = []
                for row in table.rows:
                    row_cells = []
                    for cell in row.cells:
                        cell_texts = []
                        for p in cell.paragraphs:
                            t = p.text.strip()
                            if t:
                                cell_texts.append(t)
                        if cell_texts:
                            row_cells.append(" ".join(cell_texts))
                    if row_cells:
                        # 표 한 줄을 " | "로 결합 (markdown 표처럼)
                        table_lines.append(" | ".join(row_cells))
                
                if table_lines:
                    parts.append("\n".join(table_lines))
    except Exception as e:
        # element 순회 실패 시 안전 폴백 — 기존 방식
        for p in doc.paragraphs:
            text = p.text.strip()
            if text:
                parts.append(text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        t = p.text.strip()
                        if t:
                            parts.append(t)
    
    return "\n\n".join(parts)


def _parse_hwp(uploaded_file) -> str:
    """HWP 파싱 — olefile 기반 본문 텍스트 추출."""
    try:
        import olefile
    except ImportError:
        return "[ERROR] olefile not installed. requirements.txt 확인."
    
    try:
        ole = olefile.OleFileIO(uploaded_file.read() if hasattr(uploaded_file, "read") else uploaded_file)
    except Exception as e:
        return f"[ERROR] HWP 읽기 실패: {e}"
    
    try:
        # HWP 본문은 보통 BodyText/Section* 스트림에 들어있음
        text_parts = []
        for stream_name in ole.listdir():
            if not stream_name:
                continue
            joined = "/".join(stream_name)
            if "BodyText/Section" in joined:
                try:
                    raw = ole.openstream(stream_name).read()
                    # HWP는 압축된 경우가 많아 zlib 시도
                    try:
                        import zlib
                        raw = zlib.decompress(raw, -15)
                    except Exception:
                        pass
                    # 한글 인코딩 — UTF-16 LE (HWP 본문 표준)
                    try:
                        decoded = raw.decode("utf-16le", errors="ignore")
                        # 제어 문자 제거
                        cleaned = "".join(
                            ch for ch in decoded
                            if ch.isprintable() or ch in "\n\t "
                        )
                        if cleaned.strip():
                            text_parts.append(cleaned)
                    except Exception:
                        pass
                except Exception:
                    continue
        
        ole.close()
        result = "\n\n".join(text_parts)
        return result if result.strip() else "[WARN] HWP에서 텍스트를 추출하지 못했습니다. 다른 형식으로 변환해 보세요."
    except Exception as e:
        return f"[ERROR] HWP 처리 실패: {e}"


def _parse_pdf(uploaded_file) -> str:
    """PDF 파싱 — pypdf 기반 페이지 텍스트 추출."""
    try:
        from pypdf import PdfReader
    except ImportError:
        return "[ERROR] pypdf not installed. requirements.txt 확인."
    
    try:
        reader = PdfReader(uploaded_file)
    except Exception as e:
        return f"[ERROR] PDF 읽기 실패: {e}"
    
    parts = []
    for i, page in enumerate(reader.pages):
        try:
            text = page.extract_text() or ""
            text = text.strip()
            if text:
                parts.append(text)
        except Exception:
            continue
    
    return "\n\n".join(parts)


def _parse_text(uploaded_file) -> str:
    """txt/md 파싱."""
    try:
        raw = uploaded_file.read()
        if isinstance(raw, bytes):
            # 인코딩 자동 감지 시도
            for enc in ("utf-8", "cp949", "euc-kr", "utf-16"):
                try:
                    return raw.decode(enc)
                except UnicodeDecodeError:
                    continue
            return raw.decode("utf-8", errors="ignore")
        return str(raw)
    except Exception as e:
        return f"[ERROR] 텍스트 읽기 실패: {e}"


# ============================================================================
# v3.0 Phase D — 회차 스토리라인 자동 추출 헬퍼
# ============================================================================

import re


# 회차 마커 패턴 — 4명의 전남편들 기획서 분석 결과 기반
_EPISODE_MARKER_PATTERNS = [
    # **설정(1화)**, **만남(5화)**, **사건2-1 (16화)** 등
    r'\*\*([^*]*?\(\s*(\d+)\s*화\s*\))[^*]*?\*\*',
    # ACT I, ACT II, ACT III
    r'\*\*\s*(ACT\s+[IVX]+)\s*\*\*',
    # 일반 텍스트의 (1화), 1화, 제1화 같은 패턴
    r'(?:^|\n)\s*(?:#+\s*)?(?:제)?\s*(\d+)\s*화[\.\s]',
]


def detect_episode_structure(text: str) -> dict:
    """기획서 텍스트에서 ACT 구조와 회차 마커를 자동 감지.
    
    Returns:
        {
            "has_episode_structure": bool,
            "act_markers": [{"name": "ACT I", "position": int}, ...],
            "episode_markers": [{"label": "설정(1화)", "ep_num": 1, "position": int}, ...],
            "max_episode": int,
        }
    """
    result = {
        "has_episode_structure": False,
        "act_markers": [],
        "episode_markers": [],
        "max_episode": 0,
    }
    
    if not text:
        return result
    
    # ACT 마커 — 줄 단위로 단독 등장하는 ACT I/II/III 패턴
    # **ACT I** 또는 ACT I 둘 다 매칭
    act_pattern = r'(?:^|\n)\s*\*{0,2}\s*(ACT\s+[IVX]+)\s*\*{0,2}\s*(?:\n|$)'
    for m in re.finditer(act_pattern, text):
        result["act_markers"].append({
            "name": m.group(1).strip(),
            "position": m.start(),
        })
    
    # 회차 마커 — 다음 패턴 모두 매칭
    ep_patterns = [
        # 마크다운 굵기 표시 있음: **... (N화)** 형식
        r'\*{2,}([^*\n]*?\(\s*(\d+)\s*화\s*\))[^*\n]*?\*{2,}',
        # 마크다운 없는 일반 텍스트: 줄 시작에 [라벨](N화) 형식
        # 한글/숫자/-/공백 1~40자 + (N화) 패턴, 단 같은 줄에 (N화)만 1번
        r'(?:^|\n)([^\n()]{1,40})\(\s*(\d+)\s*화\s*\)\s*(?:\n|$)',
    ]
    
    seen_eps = set()
    candidates = []
    
    for pattern in ep_patterns:
        for m in re.finditer(pattern, text):
            try:
                ep_num = int(m.group(2))
            except (ValueError, IndexError):
                continue
            if ep_num in seen_eps:
                continue
            
            # 라벨 정리 — 줄바꿈·여러 공백 제거
            raw_label = m.group(1).strip().strip('*').strip()
            # 줄바꿈을 단일 공백으로
            cleaned = re.sub(r'\s+', ' ', raw_label).strip()
            # 라벨이 너무 길면 마지막 부분만 (앞 섹션 헤더 잔재 제거)
            if len(cleaned) > 40:
                # "스토리라인 ACT I 설정" → "설정"만 남도록
                # ACT, 스토리라인 같은 헤더 단어 분리
                parts_split = re.split(r'\bACT\s+[IVX]+\b|스토리라인', cleaned)
                # 마지막 의미 있는 토막 사용
                meaningful = [p.strip() for p in parts_split if p.strip()]
                if meaningful:
                    cleaned = meaningful[-1]
            label = f"{cleaned}({ep_num}화)" if cleaned else f"EP{ep_num}"
            
            candidates.append({
                "label": label,
                "ep_num": ep_num,
                "position": m.start(),
            })
            seen_eps.add(ep_num)
    
    # 위치 순으로 정렬
    result["episode_markers"] = sorted(candidates, key=lambda x: x["position"])
    if result["episode_markers"]:
        result["max_episode"] = max(em["ep_num"] for em in result["episode_markers"])
    
    # 회차가 3개 이상이면 회차 구조가 있는 것으로 판정
    result["has_episode_structure"] = len(result["episode_markers"]) >= 3
    
    return result


def extract_episode_storylines(text: str) -> list:
    """기획서에서 회차별 스토리라인을 분리 추출.
    
    회차 마커 위치 사이의 텍스트를 해당 회차의 본문으로 추출.
    
    Returns:
        [{"ep_num": 1, "label": "설정(1화)", "content": "..."}, ...]
    """
    structure = detect_episode_structure(text)
    if not structure["has_episode_structure"]:
        return []
    
    episodes = []
    markers = sorted(structure["episode_markers"], key=lambda x: x["position"])
    
    for i, marker in enumerate(markers):
        start = marker["position"]
        # 다음 마커까지 또는 텍스트 끝까지
        end = markers[i + 1]["position"] if i + 1 < len(markers) else len(text)
        
        content = text[start:end].strip()
        # 마커 자체는 제외 — 마커 다음 줄부터
        lines = content.split("\n", 1)
        if len(lines) > 1:
            body = lines[1].strip()
        else:
            body = ""
        
        episodes.append({
            "ep_num": marker["ep_num"],
            "label": marker["label"],
            "content": body,
        })
    
    return episodes


if __name__ == "__main__":
    # 자가 검증
    test_text = """
ACT I

**설정(1화)**

40대 시호는 결혼을 포기한 채 평범한 직장생활을 한다.

**캐릭터 소개1-시호(2화)**

29세 유빈의 몸으로 깨어난다.

ACT II

**사건1(10화)**

첫 번째 전남편이 등장한다.
"""
    structure = detect_episode_structure(test_text)
    print(f"ACT 마커: {structure['act_markers']}")
    print(f"회차 마커: {structure['episode_markers']}")
    print(f"최대 회차: {structure['max_episode']}")
    print(f"회차 구조 있음: {structure['has_episode_structure']}")
    
    episodes = extract_episode_storylines(test_text)
    print(f"\n추출된 회차 본문 {len(episodes)}개:")
    for ep in episodes:
        print(f"  EP{ep['ep_num']} [{ep['label']}]: {ep['content'][:50]}...")
