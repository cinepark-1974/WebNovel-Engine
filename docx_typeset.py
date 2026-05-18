"""
docx_typeset.py — 웹소설 조판 양식 출력 모듈
==============================================
함초롬바탕 폰트, 모바일 리더 사이즈, 양쪽 정렬, 첫 줄 들여쓰기.

용도:
- 25화·50화 같은 마일스톤 시점에 일괄 조판 출력
- 작가가 모바일에서 사전 검토용으로 읽거나, 출판사 제출용으로 사용

주의:
- 매 회차마다 자동 조판하지 않음 (집필 자유 보장)
- 작가가 명시적으로 [📚 조판 다운로드] 버튼 클릭 시에만 생성

조판 양식 사양 (조판_양식_워드.docx 분석 결과):
- 페이지: 76mm × 116mm (모바일 리더 사이즈)
- 여백: 약 1mm (좁은 여백)
- 폰트: 함초롬바탕 (한국 출판 표준)
- 본문: 10.5pt
- 줄간격: 1.4 (모바일 가독성)
- 첫 줄 들여쓰기: 한 글자 (3.5mm)
- 정렬: 양쪽 정렬

변경 이력:
- v3.0: build_safe_season_docx 추가 (정수 키 정렬 보장)
- v3.0+: build_styled_episode_docx 4종 스타일 (WorkTitle/Holder/EpisodeTitle/BodyKor)
- v3.1: 5종 워드 단락 스타일 시스템 도입
        제목 / 회차제목 / 지문 / 대사 / 씬구분
        대사 = 큰따옴표 + 작은따옴표(속마음), 위아래 여백 + 넓은 줄간격
        씬구분 = *** 단독 단락, 가운데 정렬
        Word 스타일 패널에서 한 번 수정 시 같은 종류 단락 일괄 변경 가능
        + 메타정보 클리닝(author/comments 제거, 추적 정보 비움)
        + build_safe_season_docx, build_styled_episode_docx에 동일 스타일 적용
"""

from io import BytesIO
from docx import Document
from docx.shared import Pt, Twips, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ══════════════════════════════════════════════
# 조판 사양 상수
# ══════════════════════════════════════════════
TYPESET_FONT = '함초롬바탕'
TYPESET_BODY_SIZE = Pt(10.5)
TYPESET_TITLE_SIZE = Pt(11)
TYPESET_LINE_SPACING = 1.4
TYPESET_FIRST_LINE_INDENT = Pt(11)  # 약 3.5mm = 한 글자

# 모바일 리더 페이지 사이즈 (twips: 1mm ≈ 56.7 twips)
PAGE_WIDTH_TWIPS = 4082   # 약 72mm
PAGE_HEIGHT_TWIPS = 6226  # 약 110mm
MARGIN_TWIPS = 56         # 약 1mm


# ══════════════════════════════════════════════
# [v3.1] 5종 워드 단락 스타일 시스템
# ══════════════════════════════════════════════
# 워드에서 작가가 "스타일" 패널을 열어 한 번만 수정하면
# 같은 종류의 모든 단락이 일괄 변경됩니다.
# 예: 본문 글씨를 다 11pt로 키우고 싶으면 "지문" 스타일 하나만 수정.
#
# 5종:
#   제목       — 작품 제목 (16pt 굵게 가운데, 표지용)
#   회차제목   — EP1. xxx (14pt 굵게 가운데)
#   지문       — 본문 단락 (11pt 양쪽 정렬, 줄간격 1.5)
#   대사       — 따옴표 단락 (11pt 양쪽 정렬, 줄간격 1.8 + 위아래 여백)
#                큰따옴표(외부 대사) / 작은따옴표(속마음) 모두 적용
#   씬구분     — *** 단독 단락 (가운데 정렬, 위아래 여백 크게)
# ══════════════════════════════════════════════
STYLE_TITLE = '제목'
STYLE_EP_TITLE = '회차제목'
STYLE_NARRATION = '지문'
STYLE_DIALOGUE = '대사'
STYLE_SCENE_BREAK = '씬구분'

# 작품 인장 보조 스타일 (표지 IP 홀더용)
STYLE_HOLDER = 'IP홀더'


def _ensure_paragraph_styles(doc, body_font: str = '맑은 고딕'):
    """문서에 5종 단락 스타일을 정의합니다.
    
    이미 같은 이름의 스타일이 있으면 건너뜁니다.
    워드에서 사용자 정의 스타일로 인식되어 스타일 패널에 표시됩니다.
    
    Args:
        doc: Document 객체
        body_font: 본문 폰트 (기본 '맑은 고딕')
    
    Returns:
        dict: {스타일이름: Style 객체}
    """
    from docx.shared import RGBColor
    from docx.enum.style import WD_STYLE_TYPE
    
    styles = doc.styles
    existing = {s.name for s in styles}
    out = {}
    
    def _set_eastasia(style, font_name):
        """동아시아 폰트 명시 (한글 표시에 필수)."""
        rpr = style.element.get_or_add_rPr()
        rfonts = rpr.find(qn('w:rFonts'))
        if rfonts is None:
            rfonts = OxmlElement('w:rFonts')
            rpr.append(rfonts)
        rfonts.set(qn('w:eastAsia'), font_name)
        rfonts.set(qn('w:ascii'), font_name)
        rfonts.set(qn('w:hAnsi'), font_name)
    
    def _make(name, size_pt, bold=False, italic=False, color=None,
              align=None, line_spacing=None, space_before=None, space_after=None,
              first_line_indent=None):
        if name in existing:
            out[name] = styles[name]
            return
        s = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        s.font.name = body_font
        s.font.size = Pt(size_pt)
        s.font.bold = bold
        s.font.italic = italic
        if color:
            s.font.color.rgb = RGBColor(*color)
        _set_eastasia(s, body_font)
        
        pf = s.paragraph_format
        if align is not None:
            pf.alignment = align
        if line_spacing is not None:
            pf.line_spacing = line_spacing
        if space_before is not None:
            pf.space_before = space_before
        if space_after is not None:
            pf.space_after = space_after
        if first_line_indent is not None:
            pf.first_line_indent = first_line_indent
        
        out[name] = s
    
    # 제목 — 16pt 굵게 가운데
    _make(STYLE_TITLE, 16, bold=True,
          align=WD_ALIGN_PARAGRAPH.CENTER,
          line_spacing=1.3,
          space_before=Pt(0), space_after=Pt(12))
    
    # IP 홀더 — 보조 (10pt 회색 이탤릭 가운데)
    _make(STYLE_HOLDER, 10, italic=True, color=(0x66, 0x66, 0x66),
          align=WD_ALIGN_PARAGRAPH.CENTER,
          line_spacing=1.2,
          space_before=Pt(2), space_after=Pt(6))
    
    # 회차제목 — 14pt 굵게 가운데, 위 여백 크게
    _make(STYLE_EP_TITLE, 14, bold=True,
          align=WD_ALIGN_PARAGRAPH.CENTER,
          line_spacing=1.3,
          space_before=Pt(18), space_after=Pt(12))
    
    # 지문 — 본문, 양쪽 정렬, 줄간격 1.5
    _make(STYLE_NARRATION, 11,
          align=WD_ALIGN_PARAGRAPH.JUSTIFY,
          line_spacing=1.5,
          space_before=Pt(0), space_after=Pt(6))
    
    # 대사 — 양쪽 정렬, 줄간격 1.8, 위아래 여백 (한 줄 바꿈 효과)
    _make(STYLE_DIALOGUE, 11,
          align=WD_ALIGN_PARAGRAPH.JUSTIFY,
          line_spacing=1.8,
          space_before=Pt(6), space_after=Pt(6))
    
    # 씬구분 — 가운데 정렬, 위아래 여백 크게
    _make(STYLE_SCENE_BREAK, 11,
          align=WD_ALIGN_PARAGRAPH.CENTER,
          line_spacing=1.5,
          space_before=Pt(12), space_after=Pt(12))
    
    return out


def _classify_paragraph(line: str) -> str:
    """단락 한 줄을 5종 중 하나로 분류합니다.
    
    Returns:
        'scene_break' — *** 또는 ※※※ 같은 씬 구분자
        'dialogue'    — 큰따옴표 또는 작은따옴표로 시작 (외부 대사 + 속마음)
        'narration'   — 그 외 일반 본문
    
    회차제목과 작품제목은 빌더에서 별도로 결정합니다.
    """
    s = (line or "").strip()
    if not s:
        return 'narration'
    
    # 씬 구분자 — ***, ※※※, ===, ---, ◆◆◆, ★★★ 등 단독 단락
    # 글자가 같은 기호로만 3개 이상 반복되면 씬 구분
    if len(s) >= 3 and all(c in '*※=─-—◆★☆◇○●·.' for c in s):
        return 'scene_break'
    
    # 대사 — 큰따옴표 또는 작은따옴표로 시작
    # 한글 환경의 다양한 따옴표 모두 포함
    dialogue_starts = ('"', '"', '"',  # 큰따옴표 (직선/곡선)
                       "'", "'", "'",  # 작은따옴표 (직선/곡선)
                       '「', '『',      # 한국 출판 따옴표
                       '〈', '《')      # 부가
    if s.startswith(dialogue_starts):
        return 'dialogue'
    
    return 'narration'


def _add_classified_paragraph(doc, line: str, styles_map: dict):
    """분류된 스타일을 적용해서 단락을 추가합니다.
    
    Args:
        doc: Document 객체
        line: 단락 텍스트
        styles_map: _ensure_paragraph_styles의 반환값
    """
    kind = _classify_paragraph(line)
    
    if kind == 'scene_break':
        style_name = STYLE_SCENE_BREAK
    elif kind == 'dialogue':
        style_name = STYLE_DIALOGUE
    else:
        style_name = STYLE_NARRATION
    
    p = doc.add_paragraph(line.strip(), style=style_name)
    return p


def _strip_metadata(doc):
    """문서의 메타정보(코어 프로퍼티)를 비웁니다.
    
    python-docx가 기본값으로 채우는 'python-docx' 작성자, 
    'generated by python-docx' 코멘트, 2013년 생성일 등을 모두 제거합니다.
    배포 시 의도하지 않은 추적 정보가 새어 나가지 않도록 합니다.
    """
    cp = doc.core_properties
    # 비울 수 있는 문자열 필드들
    cp.author = ''
    cp.last_modified_by = ''
    cp.comments = ''
    cp.subject = ''
    cp.keywords = ''
    cp.title = ''
    cp.category = ''
    cp.content_status = ''
    cp.identifier = ''
    cp.language = ''
    cp.version = ''
    # 날짜는 비울 수 없는 필드이므로 현재 시각으로 갱신
    # (2013년 python-docx 기본 템플릿 날짜를 그대로 두지 않음)
    try:
        from datetime import datetime, timezone
        now = datetime.now(timezone.utc)
        cp.created = now
        cp.modified = now
    except Exception:
        pass


def _set_korean_font(run, font_name: str = TYPESET_FONT, size=None):
    """런(run)에 한글 폰트 + 동아시아 폰트 명시.
    
    한글 텍스트는 East Asia 폰트가 적용되도록 명시해야 함.
    """
    run.font.name = font_name
    if size is not None:
        run.font.size = size
    
    # 동아시아 폰트 명시 (한글 표시에 필수)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:cs'), font_name)


def _setup_mobile_page(doc):
    """모바일 리더 사이즈 + 좁은 여백으로 페이지 설정."""
    section = doc.sections[0]
    section.page_width = Twips(PAGE_WIDTH_TWIPS)
    section.page_height = Twips(PAGE_HEIGHT_TWIPS)
    section.top_margin = Twips(MARGIN_TWIPS)
    section.bottom_margin = Twips(MARGIN_TWIPS)
    section.left_margin = Twips(MARGIN_TWIPS * 5)   # 좌우 여백은 살짝 늘림 (가독성)
    section.right_margin = Twips(MARGIN_TWIPS * 5)


def _add_title_paragraph(doc, title_text: str):
    """제목 문단 추가 — 가운데 정렬, 굵게, 11pt."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(title_text)
    run.bold = True
    _set_korean_font(run, size=TYPESET_TITLE_SIZE)
    
    # 제목 아래 빈 줄
    doc.add_paragraph()


def _add_body_paragraph(doc, line: str, is_dialogue: bool = False):
    """본문 한 줄 추가 — 양쪽 정렬, 첫 줄 들여쓰기, 10.5pt.
    
    Args:
        doc: 문서 객체
        line: 한 줄 텍스트
        is_dialogue: 대사 여부 (대사는 들여쓰기 없음)
    """
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # 첫 줄 들여쓰기 — 대사는 없음, 서술은 있음
    pf = para.paragraph_format
    if not is_dialogue:
        pf.first_line_indent = TYPESET_FIRST_LINE_INDENT
    pf.line_spacing = TYPESET_LINE_SPACING
    pf.space_after = Pt(0)  # 문단 사이 추가 간격 없음
    pf.space_before = Pt(0)
    
    run = para.add_run(line)
    _set_korean_font(run, size=TYPESET_BODY_SIZE)


def _is_dialogue_line(line: str) -> bool:
    """대사 줄인지 판별. 따옴표·작은따옴표·낫표로 시작하면 대사."""
    return bool(line) and line.lstrip().startswith(("'", '"', '"', '"', "'", '「', '『'))


def build_safe_season_docx(
    episodes: dict,
    concept: dict,
    rating: str = "19",
    platform: str = "카카오페이지",
) -> bytes:
    """[v3.0+ 안전 빌더] session_state의 회차 본문을 그대로 사용한 통합본.
    
    기존 build_season_docx의 동기화 버그 회피용:
    - 정수 키 정렬 보장 (문자열 정렬 X)
    - 본문이 빈 회차는 경고만 + 건너뛰지 않음
    - 인장(작품 제목 + IP 홀더)은 표지에만 표시 (본문 중복 제거)
    - 회차별 페이지 분리
    
    [v3.1] 5종 워드 단락 스타일 적용 (제목/회차제목/지문/대사/씬구분)
    [v3.1] 메타정보 자동 제거
    
    동기화 버그가 의심되면 이 빌더를 우선 사용.
    """
    from docx.shared import RGBColor
    
    doc = Document()
    
    # ★ [v3.1] 5종 단락 스타일 등록
    styles_map = _ensure_paragraph_styles(doc, body_font='맑은 고딕')
    
    # A4 페이지 + 적정 여백
    section = doc.sections[0]
    section.left_margin = Pt(72)
    section.right_margin = Pt(72)
    section.top_margin = Pt(72)
    section.bottom_margin = Pt(72)
    
    # ── 표지 ──
    work_title = (concept or {}).get("title", "(제목 미지정)")
    ip_holder = (concept or {}).get("ip_holder", "블루진픽처스")
    
    # 작품 제목 (제목 스타일)
    doc.add_paragraph(work_title, style=STYLE_TITLE)
    
    # 부제 (회차 수 + 등급) — IP홀더 스타일로 대체 (작은 회색 이탤릭)
    sub_text = f"시즌 1 {len(episodes)}화 통합본 ({rating}금)"
    doc.add_paragraph(sub_text, style=STYLE_HOLDER)
    
    # IP 홀더
    doc.add_paragraph(ip_holder, style=STYLE_HOLDER)
    
    doc.add_page_break()
    
    # ── 회차별 본문 ──
    # ★ 정수 정렬 보장 (문자열 정렬 시 EP10이 EP2 앞에 오는 사고 방지)
    ep_keys_int = []
    for k in episodes.keys():
        try:
            ep_keys_int.append(int(k))
        except (ValueError, TypeError):
            continue
    ep_keys_int.sort()
    
    for ep_num in ep_keys_int:
        # ★ session_state 정수 키와 JSON 문자열 키 모두 대응
        text = episodes.get(ep_num) or episodes.get(str(ep_num), "")
        if not text or not text.strip():
            # 빈 회차 — 경고만 표시
            warn_p = doc.add_paragraph()
            warn_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            warn_run = warn_p.add_run(f"⚠️ EP{ep_num} 본문 없음")
            warn_run.font.size = Pt(11)
            warn_run.italic = True
            warn_run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
            _set_korean_font(warn_run, '맑은 고딕', Pt(11))
            doc.add_page_break()
            continue
        
        lines = text.split('\n')
        
        # 작품 인장 제거 (이미 표지에 있으니 본문 중복 제거)
        # 첫 줄이 작품 제목이면 그 이후로 시작
        idx = 0
        if lines and lines[0].strip() == work_title:
            idx = 1
            # IP 홀더도 건너뛰기
            while idx < len(lines) and not lines[idx].strip().startswith("EP"):
                idx += 1
        
        if idx >= len(lines):
            # 인장만 있고 본문 없음
            doc.add_page_break()
            continue
        
        # 회차 제목 (회차제목 스타일)
        title_line = lines[idx].strip()
        doc.add_paragraph(title_line, style=STYLE_EP_TITLE)
        
        # 본문 — 5종 스타일 자동 분류 적용
        body_lines = [l for l in lines[idx+1:] if l.strip()]
        for line in body_lines:
            _add_classified_paragraph(doc, line, styles_map)
        
        # 마지막 회차가 아니면 페이지 분리
        if ep_num != ep_keys_int[-1]:
            doc.add_page_break()
    
    # ★ [v3.1] 메타정보 클리닝
    _strip_metadata(doc)
    
    return _save_to_bytes(doc)


def build_styled_episode_docx(text: str, ep_label: str = "") -> bytes:
    """[v3.0+ / v3.1] 일반 양식 docx — 작품 인장 + 워드 5종 스타일.
    
    워드에서 일괄 수정 가능한 5종 단락 스타일을 적용:
    - 제목       — 작품 제목 (16pt 굵게 가운데, 표지용)
    - IP홀더     — IP 홀더 (10pt 회색 이탤릭)
    - 회차제목   — EP1. xxx (14pt 굵게 가운데)
    - 지문       — 일반 본문 (11pt 양쪽 정렬 줄간격 1.5)
    - 대사       — 따옴표 단락 (11pt 양쪽 정렬 줄간격 1.8 + 위아래 여백)
                  큰따옴표·작은따옴표 모두
    - 씬구분     — *** 단독 단락 (가운데 정렬 위아래 띄움)
    
    작가가 워드에서 스타일 패널을 열어 한 번 수정하면
    같은 스타일의 모든 단락이 일괄 변경됨.
    
    [v3.1] 메타정보 자동 제거 (python-docx 기본값 노출 방지)
    
    입력 텍스트 형식 (add_meta_inscription 후):
        {작품 제목}
        {IP 홀더}
        
        EP{N}. 회차 제목
        
        [본문]
    """
    doc = Document()
    
    # ★ [v3.1] 5종 단락 스타일 등록
    styles_map = _ensure_paragraph_styles(doc, body_font='맑은 고딕')
    
    # A4 페이지 + 적정 여백
    section = doc.sections[0]
    section.left_margin = Pt(72)
    section.right_margin = Pt(72)
    section.top_margin = Pt(72)
    section.bottom_margin = Pt(72)
    
    # 본문 파싱
    raw_lines = text.split('\n')
    
    # 첫 줄: 작품 제목 / 둘째 줄: IP 홀더 (인장 없으면 빈 처리)
    has_inscription = (
        len(raw_lines) >= 2
        and raw_lines[0].strip()
        and raw_lines[1].strip()
        and not raw_lines[0].strip().startswith("EP")
    )
    
    if has_inscription:
        work_title = raw_lines[0].strip()
        ip_holder = raw_lines[1].strip()
        
        # 작품 제목 + IP 홀더 (인장)
        doc.add_paragraph(work_title, style=STYLE_TITLE)
        doc.add_paragraph(ip_holder, style=STYLE_HOLDER)
        
        # 회차 제목 찾기
        idx = 2
        while idx < len(raw_lines) and not raw_lines[idx].strip():
            idx += 1
    else:
        idx = 0
    
    # 회차 제목
    if idx < len(raw_lines):
        ep_title = raw_lines[idx].strip()
        if ep_title:
            doc.add_paragraph(ep_title, style=STYLE_EP_TITLE)
            idx += 1
    
    # 본문 — 5종 스타일 자동 분류
    while idx < len(raw_lines) and not raw_lines[idx].strip():
        idx += 1
    
    body_lines = [l for l in raw_lines[idx:] if l.strip()]
    for line in body_lines:
        _add_classified_paragraph(doc, line, styles_map)
    
    # ★ [v3.1] 메타정보 클리닝
    _strip_metadata(doc)
    
    return _save_to_bytes(doc)


def build_typeset_episode(text: str, ep_label: str = "") -> bytes:
    """단일 회차를 조판 양식 docx로 변환.
    
    Args:
        text: 회차 본문 (제목 포함)
        ep_label: 회차 라벨 (예: "EP1") — docx 제목 fallback용
    
    Returns:
        DOCX 파일 bytes
    """
    doc = Document()
    _setup_mobile_page(doc)
    
    lines = [line for line in text.split('\n') if line.strip()]
    if not lines:
        return _save_to_bytes(doc)
    
    # 첫 줄을 제목으로 처리 (예: "EP1. 나는 왜 이렇게 예쁜가")
    title_line = lines[0]
    body_lines = lines[1:]
    
    _add_title_paragraph(doc, title_line)
    
    # 본문
    for line in body_lines:
        is_dlg = _is_dialogue_line(line)
        _add_body_paragraph(doc, line, is_dialogue=is_dlg)
    
    # ★ [v3.1] 메타정보 클리닝
    _strip_metadata(doc)
    
    return _save_to_bytes(doc)


def build_typeset_milestone(
    episodes: dict,
    concept: dict,
    rating: str = "19",
    milestone: int = 25,
    platform: str = "카카오페이지",
) -> bytes:
    """마일스톤 시점(25화·50화)에 1화부터 N화까지 통합 조판 docx 생성.
    
    Args:
        episodes: {ep_num: text} 회차별 본문 dict
        concept: 콘셉트 카드
        rating: "19" 또는 "15"
        milestone: 25 또는 50 (어디까지 묶을지)
        platform: 플랫폼명 (커버에 표시용)
    
    Returns:
        DOCX 파일 bytes (1~N화 통합)
    """
    doc = Document()
    _setup_mobile_page(doc)
    
    # ── 커버 페이지 ────────────────
    title = concept.get("title", "(제목 없음)")
    genre = concept.get("genre", "")
    
    # 작품 제목 (큰 글씨)
    cover_title = doc.add_paragraph()
    cover_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover_run = cover_title.add_run(title)
    cover_run.bold = True
    _set_korean_font(cover_run, size=Pt(14))
    
    # 빈 줄
    doc.add_paragraph()
    
    # 시즌 정보
    n_eps = len([k for k in episodes.keys() if str(k).isdigit() and int(k) <= milestone])
    
    info_para = doc.add_paragraph()
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_run = info_para.add_run(
        f"1화 ~ {milestone}화 모음\n{rating}금 · {genre} · {platform}"
    )
    _set_korean_font(info_run, size=Pt(10))
    
    # 페이지 나누기
    doc.add_page_break()
    
    # ── 목차 ────────────────────────
    toc_title = doc.add_paragraph()
    toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    toc_run = toc_title.add_run("목 차")
    toc_run.bold = True
    _set_korean_font(toc_run, size=Pt(12))
    
    doc.add_paragraph()
    
    # 회차 정렬
    sorted_eps = sorted(
        [k for k in episodes.keys() if str(k).isdigit() and int(k) <= milestone],
        key=lambda x: int(x)
    )
    
    for ep_key in sorted_eps:
        ep_text = episodes[ep_key]
        first_line = ep_text.split('\n')[0] if ep_text else f"EP{ep_key}"
        
        toc_item = doc.add_paragraph()
        toc_item.alignment = WD_ALIGN_PARAGRAPH.LEFT
        toc_run = toc_item.add_run(first_line)
        _set_korean_font(toc_run, size=Pt(9))
    
    doc.add_page_break()
    
    # ── 각 회차 본문 ────────────────
    for ep_key in sorted_eps:
        ep_text = episodes[ep_key]
        if not ep_text:
            continue
        
        lines = [line for line in ep_text.split('\n') if line.strip()]
        if not lines:
            continue
        
        title_line = lines[0]
        body_lines = lines[1:]
        
        # 제목
        _add_title_paragraph(doc, title_line)
        
        # 본문
        for line in body_lines:
            is_dlg = _is_dialogue_line(line)
            _add_body_paragraph(doc, line, is_dialogue=is_dlg)
        
        # 회차 끝 — 다음 회차로 페이지 나누기 (마지막 회차는 안 함)
        if ep_key != sorted_eps[-1]:
            doc.add_page_break()
    
    # ★ [v3.1] 메타정보 클리닝
    _strip_metadata(doc)
    
    return _save_to_bytes(doc)


def _save_to_bytes(doc) -> bytes:
    """docx를 bytes로 변환."""
    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


# ══════════════════════════════════════════════
# 자가 검증
# ══════════════════════════════════════════════
def _self_test():
    """모듈 자가 검증."""
    sample_text = """EP1. 테스트 회차
죽으면 끝인 줄 알았다.
그런데 깼다.
'아 씨, 이거 실화냐.'
"누구세요?"
천천히 손을 들어 올렸다.
***
손이 작았다."""
    
    # 단일 회차 (모바일 조판)
    bytes_data = build_typeset_episode(sample_text, "EP1")
    assert len(bytes_data) > 1000, "단일 회차 생성 실패"
    
    # 마일스톤 (모바일 조판)
    eps = {"1": sample_text, "2": sample_text.replace("EP1", "EP2")}
    bytes_data = build_typeset_milestone(
        eps, {"title": "테스트 작품", "genre": "현대로맨스"}, "19", 2
    )
    assert len(bytes_data) > 2000, "마일스톤 생성 실패"
    
    # ★ [v3.1] 5종 스타일 빌더 — 시즌 통합본
    season_bytes = build_safe_season_docx(
        eps, {"title": "테스트 작품", "ip_holder": "블루진픽처스"}, "19"
    )
    assert len(season_bytes) > 2000, "시즌 통합본 v3.1 생성 실패"
    
    # ★ [v3.1] 5종 스타일 빌더 — 개별 회차
    styled_bytes = build_styled_episode_docx(
        "테스트 작품\n블루진픽처스\n\n" + sample_text, "EP1"
    )
    assert len(styled_bytes) > 1000, "개별 회차 v3.1 생성 실패"
    
    # ★ [v3.1] 단락 분류 함수
    assert _classify_paragraph('"안녕하세요."') == 'dialogue', "큰따옴표 대사 분류 실패"
    assert _classify_paragraph("'속마음이다.'") == 'dialogue', "작은따옴표 속마음 분류 실패"
    assert _classify_paragraph("일반 지문입니다.") == 'narration', "지문 분류 실패"
    assert _classify_paragraph("***") == 'scene_break', "씬구분 분류 실패"
    assert _classify_paragraph("※※※") == 'scene_break', "씬구분(※) 분류 실패"
    
    # ★ [v3.1] 메타정보 클리닝 검증
    from io import BytesIO
    from docx import Document as _D
    d = _D(BytesIO(season_bytes))
    cp = d.core_properties
    assert cp.author == '', f"author 미클리닝: {cp.author!r}"
    assert cp.comments == '', f"comments 미클리닝: {cp.comments!r}"
    assert 'python-docx' not in (cp.author or ''), "python-docx 흔적 남음"
    
    return True


if __name__ == "__main__":
    print("=" * 60)
    print("docx_typeset.py — 웹소설 조판 양식 모듈")
    print("=" * 60)
    
    _self_test()
    print("\n✓ 자가 검증 통과")
    print(f"\n사양:")
    print(f"  - 폰트: {TYPESET_FONT}")
    print(f"  - 본문: {TYPESET_BODY_SIZE.pt}pt")
    print(f"  - 줄간격: {TYPESET_LINE_SPACING}")
    print(f"  - 페이지: {PAGE_WIDTH_TWIPS} × {PAGE_HEIGHT_TWIPS} twips (모바일)")
