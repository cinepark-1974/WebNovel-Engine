"""
season_report.py — 시즌 종합 검증 보고서 생성 모듈
====================================================
50화 완성 시점에 STEP 3-5에서 "🎯 시즌 종합 검증" 버튼 클릭 시 호출.

7가지 진단 항목:
1. 분량 분포 (표준편차·최단·최장)
2. 캐릭터 등장 분포 (출현 회차 수·피크 회차)
3. 정체성 키워드 반복 (A15 규칙 위반)
4. 떡밥 회수율 (개략)
5. 인접 회차 유사도
6. 회차 간 클리프행어 중복
7. 종합 일관성 점수 + 권장 후속 작업

출력: docx (4-5쪽 적정 분량)

© 2026 BLUE JEANS PICTURES
"""

import re
import statistics
from io import BytesIO
from datetime import datetime
from collections import Counter

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ══════════════════════════════════════════════
# 분석 유틸리티
# ══════════════════════════════════════════════

def _set_korean_font(run, font='맑은 고딕', size=None):
    run.font.name = font
    if size:
        run.font.size = size
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    rfonts.set(qn('w:eastAsia'), font)
    rfonts.set(qn('w:ascii'), font)


def _jaccard(a, b, n=5):
    """n-gram 자카드 유사도"""
    if not a or not b or len(a) < n or len(b) < n:
        return 0
    s_a = set(a[i:i+n] for i in range(len(a)-n+1))
    s_b = set(b[i:i+n] for i in range(len(b)-n+1))
    if not s_a or not s_b:
        return 0
    return len(s_a & s_b) / len(s_a | s_b)


def _detect_identity_keywords(text):
    """정체성 키워드 자동 감지 (나이·기간 형태)"""
    patterns = [
        (r'\d{1,3}세', "나이(N세)"),
        (r'\d{1,3}년', "기간(N년)"),
        (r'\d{1,3}살', "나이(N살)"),
        (r'마흔[일이삼사오육칠팔구]?', "한자나이(마흔X)"),
        (r'서른[일이삼사오육칠팔구]?', "한자나이(서른X)"),
        (r'쉰[일이삼사오육칠팔구]?', "한자나이(쉰X)"),
    ]
    
    matches = Counter()
    for pattern, label in patterns:
        for m in re.findall(pattern, text):
            matches[m] += 1
    
    return matches


def analyze_season(episodes_19, episodes_15, character_bible=None, plant_map=None):
    """50화 일괄 분석 → 분석 결과 dict 반환"""
    if not episodes_19:
        return None
    
    # ───── 1. 분량 분포 ─────
    lens_19 = [(int(k), len(v)) for k, v in episodes_19.items() if isinstance(v, str)]
    lens_19.sort()
    
    total_19 = sum(l for _, l in lens_19)
    total_15 = sum(len(v) for v in episodes_15.values() if isinstance(v, str)) if episodes_15 else 0
    n_eps = len(lens_19)
    avg_19 = total_19 // n_eps if n_eps > 0 else 0
    avg_15 = total_15 // n_eps if n_eps > 0 else 0
    
    shortest = min(lens_19, key=lambda x: x[1]) if lens_19 else (0, 0)
    longest = max(lens_19, key=lambda x: x[1]) if lens_19 else (0, 0)
    stdev = statistics.stdev([l for _, l in lens_19]) if len(lens_19) > 1 else 0
    
    short_eps = [ep for ep, l in lens_19 if l < 4500]
    long_eps = [ep for ep, l in lens_19 if l > 7500]
    
    # ───── 2. 캐릭터 등장 분포 ─────
    characters = {}
    if character_bible:
        # 바이블에서 자동 추출
        protag = character_bible.get('protagonist', {})
        if protag.get('name'):
            characters[protag['name']] = [protag['name']]
        for li in character_bible.get('love_interests', []):
            if li.get('name'):
                characters[li['name']] = [li['name']]
        villain = character_bible.get('villain', {})
        if villain.get('name'):
            characters[villain['name']] = [villain['name']]
        for sup in character_bible.get('supporting', [])[:5]:
            if sup.get('name'):
                characters[sup['name']] = [sup['name']]
    
    char_stats = []
    for name, keys in characters.items():
        total_count = sum(sum(v.count(k) for k in keys) for v in episodes_19.values())
        appeared_eps = []
        peak_ep, peak_count = 0, 0
        for ep_num in range(1, n_eps + 1):
            k = str(ep_num)
            if k not in episodes_19:
                continue
            ep_count = sum(episodes_19[k].count(key) for key in keys)
            if ep_count > 0:
                appeared_eps.append(ep_num)
            if ep_count > peak_count:
                peak_count = ep_count
                peak_ep = ep_num
        
        char_stats.append({
            "name": name,
            "total": total_count,
            "ep_count": len(appeared_eps),
            "peak_ep": peak_ep,
            "peak_count": peak_count,
            "ratio": len(appeared_eps) / n_eps if n_eps > 0 else 0,
        })
    
    char_stats.sort(key=lambda x: -x['total'])
    
    # ───── 3. 정체성 키워드 분포 ─────
    identity_violations = []
    total_identity_count = 0
    
    for ep_num in range(1, n_eps + 1):
        k = str(ep_num)
        if k not in episodes_19:
            continue
        text = episodes_19[k]
        kw_counter = _detect_identity_keywords(text)
        # 가장 많이 등장한 키워드만 봄
        if kw_counter:
            top_word, top_count = kw_counter.most_common(1)[0]
            total_identity_count += top_count
            if top_count >= 5:
                identity_violations.append({
                    "ep": ep_num,
                    "word": top_word,
                    "count": top_count,
                })
    
    identity_violations.sort(key=lambda x: -x['count'])
    
    # ───── 4. 떡밥 회수율 ─────
    plant_total = 0
    plant_with_payoff = 0
    if plant_map and isinstance(plant_map, dict):
        plants = plant_map.get('plants', [])
        plant_total = len(plants)
        plant_with_payoff = sum(1 for p in plants if p.get('payoff_episode'))
    
    # ───── 5. 인접 회차 유사도 ─────
    high_sim_pairs = []
    for ep in range(1, n_eps):
        if str(ep) not in episodes_19 or str(ep+1) not in episodes_19:
            continue
        end_a = episodes_19[str(ep)][-400:]
        start_b = episodes_19[str(ep+1)][:400]
        sim = _jaccard(end_a, start_b)
        if sim > 0.10:
            high_sim_pairs.append((ep, ep+1, sim))
    
    # ───── 6. 회차 간 클리프행어 중복 ─────
    endings = {ep: episodes_19[str(ep)][-300:] for ep in range(1, n_eps + 1) if str(ep) in episodes_19}
    dup_pairs = []
    for ep_a in range(1, n_eps + 1):
        for ep_b in range(ep_a + 1, n_eps + 1):
            if ep_a not in endings or ep_b not in endings:
                continue
            sim = _jaccard(endings[ep_a], endings[ep_b], n=7)
            if sim > 0.20:
                dup_pairs.append((ep_a, ep_b, sim))
    
    # ───── 7. 일관성 점수 ─────
    score_length = max(0, 100 - (len(short_eps) + len(long_eps)) * 5)
    score_identity = max(0, 100 - len(identity_violations) * 2)
    score_dup = max(0, 100 - len(dup_pairs) * 5)
    score_overall = (score_length + score_identity + score_dup) // 3
    
    # ───── 8. 일괄 독서 시뮬레이션 (5화·10화 연속 읽기 체감) ─────
    # 네이버 시리즈/카카오페이지 일괄 공개 시 독자는 5~10화를 연달아 읽음
    # 그 구간에서 반복·피로 패턴이 있는지 검증
    binge_5_warnings = []  # 5화 연속 구간에서 47 키워드 중복 감지
    binge_10_warnings = []
    
    # 5화 슬라이딩 윈도우
    for start_ep in range(1, n_eps - 4):
        window_eps = list(range(start_ep, start_ep + 5))
        # 이 5화 구간에 정체성 키워드가 얼마나 등장하는지
        total_in_window = 0
        same_words = Counter()
        for ep_num in window_eps:
            k = str(ep_num)
            if k not in episodes_19:
                continue
            kw = _detect_identity_keywords(episodes_19[k])
            if kw:
                top_word, top_count = kw.most_common(1)[0]
                total_in_window += top_count
                same_words[top_word] += top_count
        
        # 5화 구간에 동일 키워드가 25회 이상이면 경고 (회당 5회 평균)
        if total_in_window >= 25 and same_words:
            top_word, top_count = same_words.most_common(1)[0]
            binge_5_warnings.append({
                "start_ep": start_ep,
                "end_ep": start_ep + 4,
                "word": top_word,
                "count": top_count,
            })
    
    # 10화 슬라이딩 윈도우 (캐릭터 등장 균형)
    char_balance_warnings = []
    if char_stats:
        # 주요 인물 (출현 비율 50% 이상)을 기준으로 10화 구간 검사
        major_chars = [c for c in char_stats if c['ratio'] >= 0.5]
        for start_ep in range(1, n_eps - 9):
            window_eps = list(range(start_ep, start_ep + 10))
            for char in major_chars:
                # 이 10화 구간에서 0번 등장한 주요 인물이 있는가?
                appeared = False
                for ep_num in window_eps:
                    k = str(ep_num)
                    if k in episodes_19 and char['name'] in episodes_19[k]:
                        appeared = True
                        break
                if not appeared and char['name']:
                    char_balance_warnings.append({
                        "start_ep": start_ep,
                        "end_ep": start_ep + 9,
                        "missing_char": char['name'],
                    })
    
    # 중복 제거 (같은 인물이 여러 윈도우에서 누락되어도 기록은 한 번만)
    seen_chars = set()
    char_balance_unique = []
    for w in char_balance_warnings:
        key = w['missing_char']
        if key not in seen_chars:
            char_balance_unique.append(w)
            seen_chars.add(key)
    
    # 점수 보정 — 일괄 독서 시뮬레이션 점수
    score_binge = 100
    if binge_5_warnings:
        score_binge -= min(50, len(binge_5_warnings) * 3)
    if char_balance_unique:
        score_binge -= min(30, len(char_balance_unique) * 10)
    score_binge = max(0, score_binge)
    
    # 종합 점수에 binge 반영
    score_overall = (score_length + score_identity + score_dup + score_binge) // 4
    
    return {
        "n_eps": n_eps,
        "length": {
            "total_19": total_19,
            "total_15": total_15,
            "avg_19": avg_19,
            "avg_15": avg_15,
            "shortest": shortest,
            "longest": longest,
            "stdev": stdev,
            "short_eps": short_eps,
            "long_eps": long_eps,
        },
        "characters": char_stats,
        "identity": {
            "total": total_identity_count,
            "violations": identity_violations,
            "violation_count": len(identity_violations),
        },
        "plants": {
            "total": plant_total,
            "with_payoff": plant_with_payoff,
        },
        "high_sim_pairs": high_sim_pairs,
        "dup_pairs": dup_pairs,
        "binge": {
            "five_chapter_warnings": binge_5_warnings,
            "char_balance_warnings": char_balance_unique,
        },
        "scores": {
            "length": score_length,
            "identity": score_identity,
            "duplication": score_dup,
            "binge": score_binge,
            "overall": score_overall,
        },
    }


# ══════════════════════════════════════════════
# 보고서 docx 빌더
# ══════════════════════════════════════════════

def build_season_report_docx(analysis, work_title, ip_holder="블루진픽처스"):
    """분석 결과 → docx 보고서 (4-5쪽)"""
    if not analysis:
        return None
    
    doc = Document()
    
    section = doc.sections[0]
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    
    def heading(text, level=1):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        if level == 0:
            run.font.size = Pt(18)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif level == 1:
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0x19, 0x19, 0x70)
        else:
            run.font.size = Pt(12)
        _set_korean_font(run, '맑은 고딕', run.font.size)
        return p
    
    def para(text, bold=False, size=10.5, color=None, italic=False, indent=False):
        p = doc.add_paragraph()
        if indent:
            p.paragraph_format.left_indent = Inches(0.2)
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        if color:
            run.font.color.rgb = RGBColor(*color)
        _set_korean_font(run, '맑은 고딕', Pt(size))
        pf = p.paragraph_format
        pf.line_spacing = 1.4
        pf.space_after = Pt(3)
        return p
    
    def kv(key, value):
        p = doc.add_paragraph()
        key_run = p.add_run(f"  {key}  ")
        key_run.bold = True
        _set_korean_font(key_run, '맑은 고딕', Pt(10))
        val_run = p.add_run(str(value))
        _set_korean_font(val_run, '맑은 고딕', Pt(10))
        p.paragraph_format.space_after = Pt(2)
    
    # ═══ 표지 ═══
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run(work_title)
    title_run.bold = True
    title_run.font.size = Pt(16)
    _set_korean_font(title_run, '맑은 고딕', Pt(16))
    
    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_p.add_run(f"시즌 1 {analysis['n_eps']}화 — 일괄 퍼블리싱 검증 보고서")
    sub_run.bold = True
    sub_run.font.size = Pt(13)
    sub_run.font.color.rgb = RGBColor(0x19, 0x19, 0x70)
    _set_korean_font(sub_run, '맑은 고딕', Pt(13))
    
    # 사용 컨텍스트 부제
    ctx_p = doc.add_paragraph()
    ctx_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ctx_run = ctx_p.add_run("(네이버 시리즈·카카오페이지 등 일괄 공개 플랫폼 입점 점검용)")
    ctx_run.font.size = Pt(9.5)
    ctx_run.italic = True
    ctx_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    _set_korean_font(ctx_run, '맑은 고딕', Pt(9.5))
    
    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_p.add_run(f"{datetime.now().strftime('%Y년 %m월 %d일')} · {ip_holder}")
    date_run.font.size = Pt(10)
    date_run.italic = True
    date_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    _set_korean_font(date_run, '맑은 고딕', Pt(10))
    
    doc.add_paragraph()
    
    # ═══ 종합 요약 ═══
    heading("종합 요약", 1)
    
    s = analysis['scores']
    
    # 동적 평가 코멘트 (일괄 퍼블리싱 관점)
    overall = s['overall']
    if overall >= 90:
        eval_comment = "매우 우수 — 일괄 퍼블리싱 즉시 가능."
    elif overall >= 75:
        eval_comment = "우수 — 작은 보정 후 일괄 퍼블리싱 권장."
    elif overall >= 60:
        eval_comment = "양호 — 권장 후속 작업 적용 후 일괄 공개 검토."
    else:
        eval_comment = "추가 보정 필요 — 일괄 공개 전 우선순위 항목 처리 권장."
    
    para(f"독자가 5~10화를 연달아 읽는 일괄 공개 환경(네이버 시리즈·카카오페이지 등)에서 "
         f"체감되는 반복·균형·피로 패턴을 진단함. {eval_comment}", size=10.5)
    para("")
    
    kv("■ 분량 일관성", f"{s['length']} / 100")
    kv("■ 정체성 키워드 절제", f"{s['identity']} / 100")
    kv("■ 클리프행어 다양성", f"{s['duplication']} / 100")
    kv("■ 일괄 독서 체감", f"{s.get('binge', 100)} / 100  ★ 5~10화 연속 읽기 시뮬레이션")
    kv("■ 종합 점수", f"{s['overall']} / 100")
    
    doc.add_paragraph()
    
    # ═══ 1. 분량 분포 ═══
    heading("1. 분량 분포", 1)
    
    L = analysis['length']
    kv("19금 총 분량", f"{L['total_19']:,}자 (회당 평균 {L['avg_19']:,}자)")
    if L['total_15'] > 0:
        kv("15금 총 분량", f"{L['total_15']:,}자 (회당 평균 {L['avg_15']:,}자)")
    kv("분량 표준편차", f"{L['stdev']:.0f}자  (작을수록 일관됨)")
    kv("최단 회차", f"EP{L['shortest'][0]} ({L['shortest'][1]:,}자)")
    kv("최장 회차", f"EP{L['longest'][0]} ({L['longest'][1]:,}자)")
    
    if L['short_eps']:
        kv("⚠️ 분량 미달 회차", f"{L['short_eps']} (4,500자 미만)")
    if L['long_eps']:
        kv("ℹ️ 분량 초과 회차", f"{L['long_eps']} (7,500자 초과)")
    
    if not L['short_eps'] and not L['long_eps']:
        para("")
        para("■ 평가", bold=True, size=10.5)
        para("회차당 분량이 모두 정상 범위. 페이싱 안정적.", size=10)
    
    doc.add_paragraph()
    
    # ═══ 2. 캐릭터 등장 분포 ═══
    heading("2. 캐릭터 등장 분포", 1)
    
    if analysis['characters']:
        para(f"주요 인물 {len(analysis['characters'])}명 출현 통계 (19금 본문 기준):", size=10)
        para("")
        
        for char in analysis['characters'][:7]:
            ratio_pct = int(char['ratio'] * 100)
            
            # 평가
            if char['ratio'] >= 0.8:
                eval_str = "전 회차 등장"
            elif char['ratio'] >= 0.5:
                eval_str = "주요 인물 적정"
            elif char['ratio'] >= 0.3:
                eval_str = "조연 적정"
            else:
                eval_str = "출현 빈도 낮음 — 검토 권장"
            
            text = (f"  {char['name']:<15} {char['total']:>4}회 / "
                    f"{char['ep_count']}/{analysis['n_eps']}회차 ({ratio_pct}%) / "
                    f"피크 EP{char['peak_ep']} ({char['peak_count']}회)  — {eval_str}")
            para(text, size=10, indent=True)
        
        # 출현 빈도 낮은 인물 발견 시
        low_chars = [c for c in analysis['characters'] if 0 < c['ratio'] < 0.3]
        if low_chars:
            para("")
            para("■ 발견 — 출현 빈도 검토 권장", bold=True, size=10.5)
            names = ", ".join(c['name'] for c in low_chars)
            para(f"{names} — 출현 회차 30% 미만. 후속 시즌 또는 외전에서 보강 검토.", size=10)
    else:
        para("캐릭터 바이블 데이터 없음. STEP 1에서 바이블 생성 필요.", size=10, italic=True)
    
    doc.add_paragraph()
    
    # ═══ 3. 정체성 키워드 ═══
    heading("3. 정체성 키워드 반복 패턴", 1)
    
    I = analysis['identity']
    kv("위반 회차 (회차당 5회+)", f"{I['violation_count']} / {analysis['n_eps']}")
    
    if I['violations']:
        para("")
        para("가장 심한 회차 Top 5:", bold=True, size=10.5)
        for v in I['violations'][:5]:
            para(f"  EP{v['ep']}: '{v['word']}' {v['count']}회", size=10, indent=True)
        
        para("")
        para("■ 권장 — 후처리 보정 도구 적용", bold=True, size=10.5)
        para("회차당 4회까지 보존(정체성 강조용), 5회째부터 우회 표현으로 치환 권장. "
             "이 보정으로 LLM 강박 패턴이 자연스러운 문체로 정상화됨.", size=10)
    else:
        para("")
        para("■ 평가", bold=True, size=10.5)
        para("정체성 키워드 반복 패턴 발견되지 않음. A15 규칙 준수.", size=10)
    
    doc.add_paragraph()
    
    # ═══ 4. 떡밥 회수율 ═══
    heading("4. 떡밥 회수율", 1)
    
    P = analysis['plants']
    if P['total'] > 0:
        kv("전체 떡밥 수", f"{P['total']}개")
        kv("회수 회차 매핑됨", f"{P['with_payoff']}개 ({P['with_payoff']*100//P['total']}%)")
        
        if P['with_payoff'] < P['total']:
            para("")
            para("■ 권장", bold=True, size=10.5)
            para("일부 떡밥의 회수 회차가 추적되지 않음. STEP 2 떡밥 맵에 회수 회차 입력 시 "
                 "정량 검증 가능. 후속 작품에서는 떡밥 설계 단계에서 회수 회차 명시 권장.", size=10)
    else:
        para("떡밥 맵 데이터 없음.", size=10, italic=True)
    
    doc.add_paragraph()
    
    # ═══ 5. 클리프행어 다양성 ═══
    heading("5. 클리프행어 다양성", 1)
    
    kv("회차 간 중복 의심 쌍", f"{len(analysis['dup_pairs'])} / {analysis['n_eps']}")
    kv("인접 회차 유사도 10%+", f"{len(analysis['high_sim_pairs'])} / {analysis['n_eps']-1}")
    
    if analysis['dup_pairs']:
        para("")
        para("중복 의심 쌍:", bold=True, size=10.5)
        for a, b, sim in analysis['dup_pairs'][:5]:
            para(f"  EP{a} ↔ EP{b}: {sim:.1%}", size=10, indent=True)
    
    if not analysis['dup_pairs'] and not analysis['high_sim_pairs']:
        para("")
        para("■ 평가", bold=True, size=10.5)
        para(f"{analysis['n_eps']}화 동안 동일 클리프행어 반복 없음. 회차 분리 능력 우수.", size=10)
    
    doc.add_paragraph()
    
    # ═══ 6. 일괄 독서 시뮬레이션 ★ 일괄 퍼블리싱 핵심 ═══
    heading("6. 일괄 독서 시뮬레이션 (5~10화 연속 읽기)", 1)
    
    binge = analysis.get('binge', {})
    binge5 = binge.get('five_chapter_warnings', [])
    char_balance = binge.get('char_balance_warnings', [])
    
    para("일괄 공개 플랫폼에서는 독자가 5~10화를 한 자리에서 연달아 읽음. "
         "회당 검수에서는 자연스러운 패턴이라도 연속 읽기 시 피로·반복으로 체감될 수 있음.",
         size=10)
    para("")
    
    kv("5화 연속 정체성 키워드 과다 구간", f"{len(binge5)}개")
    kv("10화 연속 주요 인물 부재 구간", f"{len(char_balance)}개")
    
    if binge5:
        para("")
        para("■ 5화 연속 읽기 시 키워드 피로 구간 (Top 5):", bold=True, size=10.5)
        for w in binge5[:5]:
            para(f"  EP{w['start_ep']}~EP{w['end_ep']}: '{w['word']}' "
                 f"5화 합계 {w['count']}회 (회당 평균 {w['count']//5}회)", 
                 size=10, indent=True)
        para("")
        para("→ 5화 합계 25회 이상은 독자가 '같은 표현 반복'으로 체감하는 임계값. "
             "보정 도구로 이 구간을 우선 처리 권장.",
             size=10)
    
    if char_balance:
        para("")
        para("■ 10화 연속 주요 인물 미등장 구간:", bold=True, size=10.5)
        for w in char_balance[:5]:
            para(f"  EP{w['start_ep']}~EP{w['end_ep']}에서 '{w['missing_char']}' 등장 0회", 
                 size=10, indent=True)
        para("")
        para("→ 주요 인물이 10화 동안 사라지면 독자가 '서브 캐릭터 잊혔다'고 느낌. "
             "특히 일괄 공개에서는 더 두드러지므로 후속 시즌 또는 외전에서 보강 권장.",
             size=10)
    
    if not binge5 and not char_balance:
        para("")
        para("■ 평가", bold=True, size=10.5)
        para("연속 읽기 시 피로·부재 구간 발견되지 않음. 일괄 공개 적합.",
             size=10)
    
    doc.add_paragraph()
    
    # ═══ 7. 권장 후속 작업 ═══
    heading("7. 권장 후속 작업", 1)
    
    actions = []
    
    if analysis['identity']['violation_count'] > 5:
        actions.append((
            "정체성 키워드 보정",
            f"{analysis['identity']['violation_count']}개 회차에서 A15 규칙 위반. "
            "보정 도구 적용으로 자연스러운 문체로 정상화."
        ))
    
    low_chars = [c for c in analysis['characters'] if 0 < c['ratio'] < 0.3]
    if low_chars:
        names = ", ".join(c['name'] for c in low_chars)
        actions.append((
            "출현 빈도 낮은 인물 검토",
            f"{names} — 시즌 2 또는 외전 보강 검토."
        ))
    
    if P['total'] > 0 and P['with_payoff'] < P['total']:
        actions.append((
            "떡밥 회수 매핑",
            f"{P['total']}개 떡밥 중 {P['total']-P['with_payoff']}개의 회수 회차 사후 추적 필요."
        ))
    
    if analysis['length']['short_eps']:
        actions.append((
            "분량 미달 회차 보강",
            f"EP{analysis['length']['short_eps']} 재집필 또는 본문 보강."
        ))
    
    if not actions:
        para("특별한 권장 작업 없음. 출판 준비 가능 상태.", size=10.5)
    else:
        for i, (title, desc) in enumerate(actions, 1):
            p = doc.add_paragraph()
            
            prio_run = p.add_run(f"[우선순위 {i}] ")
            prio_run.bold = True
            prio_run.font.color.rgb = RGBColor(0xCC, 0x66, 0x00)
            _set_korean_font(prio_run, '맑은 고딕', Pt(10.5))
            
            title_run = p.add_run(f"{title}\n")
            title_run.bold = True
            _set_korean_font(title_run, '맑은 고딕', Pt(10.5))
            
            desc_run = p.add_run(f"     {desc}")
            _set_korean_font(desc_run, '맑은 고딕', Pt(9.5))
            
            p.paragraph_format.space_after = Pt(6)
    
    doc.add_paragraph()
    
    # ═══ 7. 결론 ═══
    heading("결론", 1)
    
    if overall >= 75:
        conclusion = (
            f"시즌 1 {analysis['n_eps']}화는 분량·캐릭터·클리프행어 면에서 "
            f"안정된 완성도를 보임. 종합 점수 {overall}/100. "
            f"권장 후속 작업 적용 시 90점대 회복 가능."
        )
    else:
        conclusion = (
            f"시즌 1 {analysis['n_eps']}화 검증 완료. 종합 점수 {overall}/100. "
            f"권장 후속 작업 우선 적용 후 재검증 권장."
        )
    
    para(conclusion, size=10.5)
    
    # bytes로 반환
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()
